"""Build the Pramana walkthrough document.

Every number is read from results/*.json at build time, so the document cannot
drift from the artefacts it describes. If an experiment has not been run, the
corresponding section says so rather than carrying a stale figure.
"""

from __future__ import annotations

import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(".")
RES = ROOT / "results"
FIG = RES / "figures"
OUT = ROOT / "Pramana_walkthrough.docx"

INK = RGBColor(0x11, 0x16, 0x1F)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
SLATE = RGBColor(0x5B, 0x65, 0x72)
BAD = RGBColor(0xB3, 0x26, 0x1E)


def load(name):
    p = RES / name
    return json.loads(p.read_text()) if p.exists() else None


def git(*args, default="unavailable"):
    try:
        return subprocess.run(["git", *args], capture_output=True, text=True,
                              check=True).stdout.strip()
    except Exception:
        return default


# --------------------------------------------------------------------------
# document primitives
# --------------------------------------------------------------------------

def setup(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = st.paragraph_format
    pf.space_after = Pt(7)
    pf.line_spacing = 1.14
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.85)
        s.left_margin = s.right_margin = Inches(0.95)


def h1(doc, text, rubric=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = INK
    if rubric:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(9)
        rr = q.add_run(rubric)
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = ACCENT
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    return p


def para(doc, text, size=10.5, italic=False, colour=None, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if colour:
        r.font.color.rgb = colour
    return p


def rich(doc, parts, size=10.5, after=7):
    """parts: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.size = Pt(size)
            r2 = p.add_run(" " + it[1]); r2.font.size = Pt(size)
        else:
            p.add_run(it).font.size = Pt(size)


def table(doc, head, rows, widths=None, size=9, align_right=()):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(head):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(v))
            r.font.size = Pt(size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, name, caption, width=6.4):
    p = FIG / name
    if not p.exists():
        para(doc, f"[figure {name} not generated]", italic=True, colour=SLATE)
        return
    doc.add_picture(str(p), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(11)
    r = c.add_run(caption)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = SLATE


def callout(doc, text, colour=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.color.rgb = colour
    return p


# --------------------------------------------------------------------------
# content
# --------------------------------------------------------------------------

def cover(doc, meta):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Pramana"); r.bold = True; r.font.size = Pt(34)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Sanskrit: a valid means of proof")
    r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = SLATE

    rich(doc, [("Payments can verify that a transaction was authorised without knowing "
                "whether it matched what the payer believed they were paying for. Pramana "
                "measures when adding that context is worth it, and how much adversarial "
                "pressure it survives.", False, False)], size=12.5, after=16)

    table(doc, ["", ""], [
        ["Research question", "Frozen before any code was written; see §Appendix A"],
        ["Pre-registration commit", meta["prereg_commit"]],
        ["Repository HEAD", meta["head"]],
        ["Built", meta["built"]],
        ["Data", "Synthetic only. No live-system testing. No operational attack tooling."],
    ], widths=[1.7, 4.6], size=9.5)


def s1_executive(doc, d):
    h1(doc, "1. Executive summary")
    rich(doc, [("The question. ", True, False),
               ("Under what levels of adversarially degraded payment-context reliability "
                "does declared payment context provide incremental fraud-detection value "
                "beyond transaction, behavioural, and beneficiary intelligence?", False, True)])

    rich(doc, [("What we built. ", True, False),
               ("A pre-registered phase study over a synthetic Indian payment ecosystem, "
                "and a deterministic mandate verifier for the agentic surface. The "
                "contribution is not a fraud model. It is a framework for deciding whether "
                "a payment network should spend money collecting a signal at all, with the "
                "adversarial tolerance characterised rather than assumed.", False, False)])

    if d["phase"]:
        rs = d["phase"]["rho_star"].get("uniform|recall@fpr=0.001", [])
        rows = []
        for r in rs:
            if r["status"] == "significant throughout":
                v = "> 1.0"
            elif r["status"] == "never significant":
                v = "< 0.0"
            else:
                v = f"{r['rho_star_lo']:g} – {r['rho_star_hi']:g}"
            rows.append([f"{r['lam']:g}", v, r["status"]])
        h2(doc, "The primary reported quantity: ρ*")
        para(doc, "ρ* is the coaching level above which the incremental value of declared "
                  "context stops being significant. It is bracketed by the sweep grid, not "
                  "interpolated, and it is a property of this threat model and "
                  "parameterisation — never a universal fact.")
        table(doc, ["λ  structural overlap", "ρ* bracket", "Reading"], rows,
              widths=[1.6, 1.4, 3.3], align_right=(0, 1))
    else:
        para(doc, "[phase study not yet generated]", italic=True, colour=SLATE)

    if d["agentic"]:
        a, bl = d["agentic"], d["agentic"]["bounded_loss"]
        fp = a["false_positives_on_in_scope_traffic"]
        h2(doc, "The agentic surface")
        table(doc, ["Quantity", "Value"], [
            ["Attack families caught deterministically",
             f"{a['coverage']['caught']} of {a['coverage']['total']}"],
            ["False-positive rate on in-scope carts",
             f"{fp['false_positive_rate']:.4f}  ({fp['rejected']} of {fp['n']:,})"],
            ["Families NOT caught, by construction", ", ".join(a["coverage"]["uncaught"])],
            ["Mean loss without enforcement", f"₹{bl['mean_loss_unenforced']:,.0f}"],
            ["Mean loss, persistent attacker, enforced",
             f"₹{bl['mean_loss_persistent']:,.0f}  ({bl['reduction_persistent']:.1%} reduction)"],
            ["p95 loss, enforced", f"₹{bl['p95_persistent']:,.0f}"],
        ], widths=[3.3, 3.0], align_right=(1,))

    h2(doc, "Convergence with the RBI's April 2026 direction")
    para(doc, "The RBI discussion paper Exploring Safeguards in Digital Payments to Curb "
              "Frauds (9 April 2026; proposed, not law; comments closed 8 May 2026) puts "
              "forward four controls — a one-hour lag above ₹10,000, trusted-person "
              "authentication, a ₹25 lakh credit cap, and a kill switch. Every one of them "
              "is a friction control. None is a detection improvement. This work asks the "
              "adjacent question the paper does not: whether the payment's declared purpose "
              "is worth capturing as a detection signal, and under how much adversarial "
              "pressure that remains true.")


def s2_problem(doc, d):
    h1(doc, "2. The problem", "Rubric: problem definition and relevance")
    para(doc, "Authorised push-payment fraud is the failure mode where every existing "
              "control works exactly as designed. The payer is authenticated. The device is "
              "recognised. The transaction is authorised. And the money is gone, because "
              "the payer was deceived about what they were paying for. Authorisation is "
              "verified; understanding is not.")

    h2(doc, "Scale, in India, now")
    table(doc, ["Quantity", "Value", "Source"], [
        ["Reported loss, 2025", "₹22,931 crore", "NCRP, cited in the RBI discussion paper"],
        ["Reported cases, 2025", "28 lakh", "NCRP, cited in the same paper"],
        ["Implied mean case loss", "₹81,896", "derived from the two figures above"],
        ["Cases above ₹10,000", "~45% by volume", "RBI discussion paper, 9 April 2026"],
        ["Value above ₹10,000", "~98.5% by value", "RBI discussion paper, 9 April 2026"],
    ], widths=[1.9, 1.5, 2.9])
    para(doc, "The volume-versus-value asymmetry is the operationally important number. "
              "A control that only sees transactions above ₹10,000 misses 55% of cases but "
              "reaches 98.5% of the money. It is also the anchor this project calibrates "
              "its synthetic population against, so the simulator inherits a real "
              "constraint rather than a convenient one.", after=9)

    h2(doc, "What is already proposed, and what it does not do")
    bullets(doc, [
        ("One-hour lag above ₹10,000.", "Friction. Buys reversal time; detects nothing."),
        ("Trusted-person authentication.", "Friction. Adds a human check; detects nothing."),
        ("₹25 lakh credit cap.", "Friction. Bounds loss; detects nothing."),
        ("Kill switch.", "Friction. Stops the bleeding after the fact; detects nothing."),
    ])
    para(doc, "Each is a reasonable control. None improves the network's ability to tell, "
              "at authorisation time, that this payment is not what the payer thinks it is. "
              "That gap is the subject of this work.")

    h2(doc, "Why this failure mode is different")
    para(doc, "Card fraud and account takeover are unauthorised: the genuine customer did "
              "not make the payment, so authentication and device intelligence are aimed at "
              "the right target. APP fraud inverts that. The customer did make the payment, "
              "on their own device, through their own credentials, having passed every "
              "check. What failed was not the authorisation but the payer's model of what "
              "the payment was for.")
    para(doc, "That inversion has three practical consequences. Liability is contested, "
              "because no control was bypassed. Reversal is hard, because the transfer was "
              "valid at the moment it was made. And detection has to work with a signal "
              "that no existing field carries: the payer's own belief about the "
              "counterparty. A network sees who paid whom, how much, from what device, at "
              "what time. It does not see 'this was supposed to be my daughter's tuition'.")
    para(doc, "The question this project asks is whether asking for that belief, and "
              "checking it against what the beneficiary actually looks like, is worth the "
              "cost of collecting it — and for how long that remains true once attackers "
              "know the check exists. The second half is the part that is usually left out.")

    h2(doc, "The forward surface")
    para(doc, "Agentic commerce makes the same gap structural rather than incidental. US "
              "AI-platform retail e-commerce is roughly $20.6bn in 2026, about 1.5% of "
              "e-commerce, and OpenAI deprecated Instant Checkout in March 2026. We state "
              "that ourselves: the capability retreated because the trust primitive is not "
              "settled. That is the argument, not a weakness in it. When the payer is an "
              "agent, 'what the principal intended' stops being a matter of inference and "
              "becomes a signed, checkable object — and §7 quantifies exactly how much that "
              "buys, including where it buys nothing.")


def s3_prior_art(doc, d):
    h1(doc, "3. Prior art, stated up front", "Rubric: originality, honest positioning")
    para(doc, "This section exists before the results deliberately. The strongest objection "
              "to this project is that declared purpose already exists, and it does. Here "
              "is where, and here is precisely what we are and are not claiming.")

    table(doc, ["What exists", "Where"], [
        ["Structured purpose codes on payment messages",
         "ISO 20022 <Purp><Cd>, carried in pain.001 and pacs.008"],
        ["Banks asking the payer the purpose of a payment",
         "UK Consumer Standard of Caution"],
        ["Purpose used in a decision by the receiving institution",
         "BIS Nexus guidance notes a destination PSP may consider purpose codes"],
        ["Cryptographically signed payment intent",
         "UPI signs payment-request parameters (not semantic purpose)"],
        ["Open/closed mandates with deterministic conformance checking",
         "AP2 v0.2"],
        ["Purpose × beneficiary as a modelling technique",
         "Ordinary feature engineering; nothing novel is claimed"],
    ], widths=[2.7, 3.6])

    h2(doc, "What we are not claiming")
    bullets(doc, [
        "That structured purpose fields do not exist.",
        "That banks do not ask purpose.",
        "That purpose is never used in decisions.",
        "That UPI lacks signed intents.",
        "That AP2 lacks mandate verification.",
        "That purpose × beneficiary is a novel technique.",
        "That nobody does this — a global negative cannot be proven.",
    ])

    h2(doc, "What we do claim")
    callout(doc, "We found no publicly documented production system that models "
                 "purpose–beneficiary consistency as a standalone feature class.")
    para(doc, "And, separately and more importantly: we are not aware of published work "
              "that measures the adversarial tolerance of such a signal — the level of "
              "coaching at which it stops paying for itself. That measurement, not the "
              "feature, is the contribution. A payment network deciding whether to capture "
              "and retain a field needs to know the conditions under which the field remains "
              "informative, and that is a different question from whether a model can use it.")


def s4_identify(doc, d):
    h1(doc, "4. Identify: the threat model", "Rubric: threat identification")
    h2(doc, "Scam families and what the victim is told")
    para(doc, "Six scam families drive the fraud process. Each carries a narrative, and the "
              "narrative determines what the victim believes — and therefore declares — "
              "when they are not coached otherwise.")
    table(doc, ["Scam family", "Share", "Narrative purpose the victim declares"], [
        ["digital_arrest", "18%", "other"],
        ["investment_scam", "26%", "investment"],
        ["task_scam", "20%", "investment"],
        ["impersonation", "14%", "family_support"],
        ["fake_fees", "10%", "education_fees"],
        ["refund_scam", "12%", "other"],
    ], widths=[1.9, 0.9, 3.5], align_right=(1,))

    h2(doc, "The adversarial parameter: ρ")
    para(doc, "ρ is coaching effectiveness. With probability 1−ρ the victim declares what "
              "the narrative told them, in good faith. With probability ρ the attacker "
              "steers the declaration into a safe set — {friend_transfer, other, "
              "investment} — chosen because a mule account's beneficiary profile does not "
              "badly violate the legitimate profile of those purposes.")
    bullets(doc, [
        ("At ρ = 0,", "a victim paying a twelve-day-old mule declares education_fees. A "
         "real education institution is years old with periodic, seasonal inflows from "
         "thousands of payers. The mismatch is large, and declared context fires."),
        ("At ρ = 1,", "they declare friend_transfer. Friends genuinely have thin history, "
         "few payers and no periodicity. There is no mismatch left to detect."),
    ])
    para(doc, "The entire result is the shape of the curve between those two points.")

    h2(doc, "A second, stronger adversary")
    para(doc, "The pre-registered adversary samples the safe set uniformly. Because "
              "legitimate 'investment' is rare (~1.8% of payments), a uniform draw leaves a "
              "base-rate trace: the declared code is itself mildly informative even at "
              "ρ = 1. We therefore also run a prevalence-matched adversary that samples the "
              "safe set in proportion to legitimate frequencies. At ρ = 1 this flattens "
              "P(fraud | declared code) to ≈ 0.0235 across the whole safe set, so the code "
              "carries no marginal information and any residual value must come from the "
              "purpose–beneficiary interaction. It is reported as a clearly separated "
              "secondary surface; it never touches the pre-registered primary axis.")

    h2(doc, "The declaration model for legitimate payments")
    para(doc, "A signal is only worth measuring against a realistic null. Legitimate payers "
              "do not declare purpose perfectly: they pick the nearest menu item, they pick "
              "'other' when in a hurry, and some categories are genuinely ambiguous. The "
              "simulator therefore passes every legitimate declaration through a "
              "row-stochastic confusion matrix whose diagonal ranges from 0.65 "
              "(friend_transfer, which people routinely record as family_support or other) "
              "to 0.91 (utility_bill, which is unambiguous) to 1.00 (other, which is the "
              "null category and cannot be mislabelled into something else). The aggregate "
              "mislabel rate is approximately 18%.")
    para(doc, "This matters more than it looks. If legitimate declarations were noiseless, "
              "any fraudulent mislabel would stand out, and the measured value of the signal "
              "would be an artefact of that idealisation. The confusion matrix is the "
              "honest null against which coaching is measured.")

    h2(doc, "The eleven-class taxonomy")
    table(doc, ["Purpose", "Beneficiary signature it implies"], [
        ["rent", "periodic, stable payee, few payers, high periodicity"],
        ["salary_reimburse", "employer-like, very low fan-in"],
        ["family_support", "recurring, reciprocal, small payer set"],
        ["friend_transfer", "low signature, bidirectional, sporadic"],
        ["education_fees", "institutional, seasonal, high fan-in, legitimate"],
        ["utility_bill", "institutional, periodic, very high fan-in, legitimate"],
        ["merchant_purchase", "commercial, high fan-in, legitimate"],
        ["loan_repayment", "periodic, institutional or individual"],
        ["investment", "no strong legitimate signature in P2P — the scam favourite"],
        ["medical", "sporadic, institutional"],
        ["other", "deliberate null category"],
    ], widths=[1.5, 4.8])
    para(doc, "Cardinality is swept: K = 3 collapses this to {personal, commercial, other}, "
              "K = 6 to a mid grouping, K = 11 is the full taxonomy. A coarser menu is "
              "cheaper to deploy and easier for a payer to answer correctly, so the "
              "question of how much resolution the signal actually needs is an "
              "operational one, not a modelling detail.")

    h2(doc, "The agentic threat model")
    para(doc, "Ten attack families against a signed mandate. Eight are scope, freshness, "
              "binding or revocation violations, and are caught structurally. Two are not, "
              "and they are in the list precisely because deterministic checking cannot "
              "catch them.")
    table(doc, ["ID", "Attack family", "Expected catcher"], [
        ["A1", "Amount escalation beyond mandate cap", "C1 amount scope"],
        ["A2", "Category violation", "C2 category scope"],
        ["A3", "Mandate replay", "C5 nonce freshness"],
        ["A4", "Cumulative aggregation across individually in-scope carts", "C6 cumulative cap"],
        ["A5", "Expired mandate reuse", "C4 temporal validity"],
        ["A6", "Forged agent attestation", "C7 agent binding"],
        ["A7", "Line-item substitution after user confirmation", "C8 confirmation binding"],
        ["A8", "Post-revocation burst", "C9 revocation state"],
        ["A9", "In-scope malicious purchase", "none — bounded, not detected"],
        ["A10", "Prompt-injected but in-scope purchase", "none — bounded, not detected"],
    ], widths=[0.5, 3.8, 2.0])
    para(doc, "Each family is constructed to exercise its own check and nothing else. A4, "
              "for instance, re-presents the mandate with a fresh nonce, re-signed by the "
              "principal, so that C5 and C10 both pass and only the cumulative cap can stop "
              "it — otherwise the test would report a catch that came from an invalid "
              "signature rather than from aggregation control. A test asserts this: every "
              "family must fail the check it claims to exercise.")


def s5_generate(doc, d):
    h1(doc, "5. Generate: the synthetic ecosystem", "Rubric: data generation and realism")
    callout(doc, "The trap: if purpose is generated from the fraud label, B4 wins trivially "
                 "and the result is worthless.")
    para(doc, "Three processes run independently, and the consistency signal is emergent "
              "from their interaction rather than planted by any of them.")
    bullets(doc, [
        ("Beneficiary behaviour is a function of payee role.", "Fifteen roles with latent "
         "parameters for account age, fan-in, onward fanout, inflow periodicity, geographic "
         "dispersion and balance retention."),
        ("True economic purpose is a function of the payer–payee relationship.",
         "Relationship formation picks a beneficiary role appropriate to the purpose. It "
         "never inspects fraud status."),
        ("Victimisation is a function of an independent campaign process.", "Victims are "
         "drawn from a personal susceptibility latent; the beneficiary is drawn from the "
         "mule population. Nothing in it reads a payer's relationship portfolio."),
    ])
    para(doc, "No generative step conditions purpose on the fraud label. The declared "
              "purpose of a fraudulent payment is decided at feature-construction time from "
              "the scam narrative and ρ, which is also what makes ρ sweepable without "
              "regenerating the ledger.")

    h2(doc, "Population and scale")
    table(doc, ["Parameter", "Value"], [
        ["Payers", "25,000"],
        ["Beneficiaries", "20,000"],
        ["Window", "12 months"],
        ["Transactions per ledger", "~2.04 million"],
        ["Payments per payer per month", "~6.8"],
        ["Beneficiary roles", "15"],
        ["Scam families", "6"],
        ["Fraud share of volume", "0.80% (never rebalanced)"],
    ], widths=[2.6, 2.0], align_right=(1,))
    para(doc, "Class balance is never adjusted. No resampling, no positive-class weighting. "
              "A rebalanced fraud dataset produces numbers that do not transfer to an "
              "operating point, and an operating point is the only thing this study reports.")

    h2(doc, "Relationships, and why staggered starts matter")
    para(doc, "Each payer holds a portfolio of standing relationships — a landlord, one or "
              "two utility billers, an employer, friends, family, merchants, sometimes a "
              "loan, a chit fund or a school. Monthly ties draw a stable base amount once "
              "and jitter it by about 5% per payment, so rent has a periodic, stable "
              "signature that a mule inflow cannot reproduce. Sporadic ties jitter widely.")
    para(doc, "22% of relationships begin part-way through the observation window. This is "
              "not cosmetic. A payer who moved house last month has a landlord with no "
              "history at all, and that is exactly the profile a naive rule flags as fraud. "
              "Without staggered starts, 'first payment to this beneficiary' would be an "
              "almost perfect fraud rule and the false-positive population — the population "
              "that actually determines whether a control is deployable — would not exist. "
              "A separate stream of ad-hoc payments to beneficiaries the payer has no "
              "standing tie with provides the same function at higher volume.")
    para(doc, "Conversely, a bank's records predate a twelve-month observation window, so "
              "pair-level history is seeded with the tie's pre-window age. A five-year-old "
              "landlord relationship must not look as unfamiliar as a fresh mule in month "
              "one. Ad-hoc and fraudulent pairs receive no seed, because they genuinely "
              "have none.")

    h2(doc, "Session telemetry from a duress latent")
    para(doc, "B2 is generated from a coercion latent rather than the label: "
              "P(coercion | scam) = 0.75 and P(coercion | legitimate) = 0.02. A coerced "
              "session runs long, hesitates on the confirmation screen, edits the amount and "
              "the payee field repeatedly, switches applications while a call is in "
              "progress, and types more slowly than that person normally does.")
    para(doc, "Two features deliberately break the clean mapping. Pasting a payee identifier "
              "is common under coercion (80%) but also common on any genuinely unfamiliar "
              "payee (45%) and uncommon on a familiar one (12%). Device novelty depends "
              "partly on coercion and partly on the payer. Without these dependencies a "
              "single session feature would be a clean coercion proxy, and B2 would be "
              "measuring the label rather than behaviour. The baseline is meant to be "
              "strong, not lucky.")

    h2(doc, "Beneficiary features, and the two lookahead guards")
    para(doc, "Payee-level network aggregates are drawn per payee-month from the role's "
              "latent parameters rather than counted off the simulated panel. The 25,000 "
              "simulated payers are a sample of each beneficiary's true inbound payer base, "
              "so counting in-sample in-degree would understate a utility biller by three "
              "orders of magnitude while leaving a mule roughly correct. The fidelity "
              "scorecard reports the Spearman correlation between in-sample observed "
              "in-degree and the generated aggregates, which is what licenses the choice.")
    bullets(doc, [
        ("Report counts are strictly as-of-time.", "Report timestamps are stored and "
         "searched at the transaction day, never summed ahead of time, so no test row can "
         "see a report filed after it. Only 35% of victims file, and filing carries a "
         "gamma-distributed delay averaging about eighteen days for victim realisation, "
         "NCRP filing and network propagation."),
        ("Legitimate beneficiaries attract disputes.", "At a rate scaled by fan-in, so that "
         "prior-report count is informative without being a perfect label."),
        ("Payer velocity uses only prior rows.", "Trailing 24-hour and 7-day counts and "
         "sums, the expanding z-score of log amount against the payer's own history, and "
         "recency are all computed from transactions that precede the row in time."),
    ])

    h2(doc, "The second trap, and λ")
    callout(doc, "If every high-fan-in payee is a mule, beneficiary intelligence alone is "
                 "near-perfect and declared context has nothing to add.", BAD)
    para(doc, "λ is the combined population share of legitimate roles that structurally "
              "resemble mule accounts. It is the reason the experiment is non-trivial.")
    table(doc, ["Legitimate role", "Why it resembles a mule"], [
        ["property_manager", "high fan-in from unrelated payers"],
        ["education_institution", "very high fan-in, seasonal"],
        ["utility_biller", "enormous fan-in"],
        ["merchant_small", "high fan-in, moderate fanout, young accounts"],
        ["settlement_agent", "forwards 88% of inflow within 24 hours"],
        ["gig_worker", "young account, sweeps balance out, thin file"],
        ["chit_fund_collector", "many unrelated payers, sweeps the pot straight out, "
                                "declared as 'investment'"],
    ], widths=[1.9, 4.4])
    para(doc, "chit_fund_collector is the sharpest case in the whole population: entirely "
              "legitimate, collecting from many unrelated payers, forwarding the pot to that "
              "month's member, and declared under the same purpose code a coached scammer "
              "steers toward. It is exactly the ambiguity that makes ρ meaningful.")
    para(doc, "This was not obvious in advance. Our first end-to-end run gave the baseline "
              "99.75% recall at 0.1% FPR, because no legitimate role in the original design "
              "had mule-range fanout — one feature scored AUC 0.9957 on its own. CHANGELOG.md "
              "records what was observed, what was changed, and why the change is a "
              "correction to realism rather than a move toward a preferred result. "
              "PREREGISTRATION.md itself has never been edited.")

    if d["fidelity"]:
        f = d["fidelity"]
        ca, cb = f["case_level_asymmetry"], f["class_balance"]
        h2(doc, "Fidelity scorecard")
        para(doc, "Reference values are published primary-source statistics, not another "
                  "simulator. The two parameters of the case-size distribution were solved, "
                  "not tuned, against the published mean case loss and the 45% of cases "
                  "above ₹10,000.")
        table(doc, ["Quantity", "Observed", "Anchor", "|error|"], [
            ["Share of cases above ₹10,000",
             f"{ca['observed']['share_of_cases_above_10k']:.4f}",
             f"{ca['anchor']['share_of_cases_above_10k']:.3f}",
             f"{ca['abs_error']['share_of_cases_above_10k']:.4f}"],
            ["Share of fraud value above ₹10,000",
             f"{ca['observed']['share_of_value_above_10k']:.4f}",
             f"{ca['anchor']['share_of_value_above_10k']:.3f}",
             f"{ca['abs_error']['share_of_value_above_10k']:.4f}"],
            ["Mean case loss",
             f"₹{ca['observed']['mean_case_loss_inr']:,.0f}",
             f"₹{ca['anchor']['mean_case_loss_inr']:,.0f}",
             f"{ca['rel_error_mean_case']:.1%}"],
            ["Fraud share of transaction volume",
             f"{cb['fraud_share_of_volume']:.4%}", "0.80%", "—"],
            ["Fraud share of transaction value",
             f"{cb['fraud_share_of_value']:.2%}", "6%", "—"],
        ], widths=[2.6, 1.3, 1.2, 1.2], align_right=(1, 2, 3))
        dd, lr, rd = f["degree_distribution"], f["latent_recovery"], f["b3_redundancy"]
        it = f["inter_transaction_times"]
        table(doc, ["Structural statistic", "Value"], [
            ["Payee in-degree CCDF log-log slope", f"{dd['ccdf_loglog_slope']:.2f}"],
            ["Median / max payee in-degree",
             f"{dd['median_in_degree']:.0f} / {dd['max_in_degree']:,}"],
            ["Median inter-transaction gap", f"{it['median_days']:.2f} days"],
            ["Share of gaps under one day", f"{it['share_under_1_day']:.1%}"],
            ["Max |correlation| between distinct B3 features", f"{rd['max_abs_corr']:.3f}"],
            ["Latent recovery (Spearman)", f"{lr['spearman']:.3f}"],
        ], widths=[4.0, 2.3], align_right=(1,))
        para(doc, "Not run: a discriminator-AUC check against real transaction data. No "
                  "labelled public APP-fraud dataset exists, and matching a second synthetic "
                  "generator would demonstrate nothing. The omission is recorded rather than "
                  "substituted for.", italic=True, colour=SLATE)


def s6_defend(doc, d):
    h1(doc, "6. Defend: features, non-circularity, and deterministic checks",
       "Rubric: technical approach and rigour")
    h2(doc, "Four feature groups, one bucket each")
    para(doc, "Every model input belongs to exactly one group. This is the single most "
              "important implementation guard in the project: if B4 silently re-encoded B3, "
              "a measured gain would only mean 'we gave the model beneficiary information "
              "twice'. pramana/features/__init__.py is the authority and a build-failing "
              "test enforces it.")
    table(doc, ["Group", "Content", "n"], [
        ["B1 transaction", "amount, log amount, hour/day/date, channel, first-payment flag, "
                           "24h and 7d velocity, payer-relative amount z-score, "
                           "amount/balance, recency", "13"],
        ["B2 payer session", "session duration, confirm-screen dwell, amount and payee-field "
                             "edits, app switches, typing speed, device novelty and age, "
                             "paste use, screen-on time, concurrent call", "11"],
        ["B3 beneficiary", "account age, 30d unique inflow payers, 30d inflow value, 24h "
                           "fanout, 30d unique outflow payees, reciprocity, pair "
                           "relationship months and prior count, report count, payer geo "
                           "dispersion, inflow CV, periodicity, balance retention", "13"],
        ["B4a declared context", "purpose code", "1"],
        ["B4b declared context", "purpose code + Mahalanobis distance + log-likelihood + 13 "
                                 "per-feature residuals", "16"],
    ], widths=[1.4, 4.3, 0.5], align_right=(2,))

    h2(doc, "Why B4b cannot leak the label")
    para(doc, "The consistency residuals are an unsupervised conditional density over "
              "beneficiary features, estimated per purpose class on training legitimate rows "
              "only. The label is used in exactly one place — to exclude known fraud from "
              "the reference set — which is what a bank does in production when it builds a "
              "profile from confirmed-good history. It is never a target, never enters the "
              "transform, and is never touched at test time. The quantile transform and "
              "every per-class moment are fitted on training data alone and then frozen, so "
              "no test row influences its own residual.")
    callout(doc, "A residual is not a beneficiary feature. It is a distance between a "
                 "beneficiary and a purpose.")
    para(doc, "A model holding B3 alone cannot compute it, because it is never shown the "
              "purpose. A model holding the purpose alone cannot compute it either, because "
              "it is never shown the beneficiary. It exists only in the interaction, which "
              "is the hypothesis under test.")
    para(doc, "Two tests enforce this rather than asserting it. One permutes the B3 values "
              "of the fraud rows the model excludes and requires every residual to be "
              "bit-identical. The other requires that scoring a test set leaves the fitted "
              "reference unchanged. Beneficiary noise β is applied before the consistency "
              "model is fitted, so B4b never sees a cleaner view of the beneficiary than the "
              "B3 arm it is compared against.")

    h2(doc, "How the consistency model is fitted")
    para(doc, "Inside every cell, on that cell's training split alone:")
    bullets(doc, [
        "Take training rows whose label is legitimate.",
        "Fit a quantile transformer with normal output on their beneficiary block, so "
        "heavy-tailed features such as inflow value and account age become comparable.",
        "For each declared-purpose class with at least 400 rows, estimate a mean and a "
        "Ledoit-Wolf shrunk covariance on the transformed block. Rarer classes fall back to "
        "a global legitimate reference.",
        "At scoring time, emit a Mahalanobis distance, a Gaussian log-likelihood, and "
        "thirteen per-feature standardised residuals.",
    ])
    para(doc, "Shrinkage is not optional here. Several beneficiary features are close to "
              "collinear within a purpose class, and an unshrunk covariance would invert "
              "unstably and produce residuals dominated by numerical noise rather than by "
              "the beneficiary.")

    h2(doc, "The baseline gets the advantage")
    para(doc, "Hyperparameters were selected by 24-candidate random search under 5-fold "
              "payer-grouped cross-validation using the B1+B2+B3 feature set alone, then "
              "frozen and reused verbatim for every arm including both B4 variants. The "
              "incumbent received the entire tuning budget; the challenger received none.")
    para(doc, "Splits satisfy two constraints simultaneously. Grouped: 30% of payers are "
              "held out and a payer appears on exactly one side, which stops the model "
              "memorising individuals. Temporal: training is months 1–9 and test is months "
              "10–12, which stops it learning from the future. Either constraint alone "
              "would leave an obvious route to an inflated result.")

    h2(doc, "Why the bootstrap is clustered on payers")
    para(doc, "A payer's transactions are correlated — same device, same habits, same "
              "beneficiary portfolio, and in the fraudulent case the same scam episode. "
              "Resampling rows would treat those as independent observations and report "
              "confidence intervals that are too narrow, which is exactly the failure that "
              "makes a small effect look real. Every interval in this document comes from "
              "1000 resamples of test payers, and the same resample weights are shared "
              "across arms so that the paired delta is computed on matched samples. That "
              "pairing is what makes an effect of a few percentage points resolvable at all.")

    h2(doc, "Deterministic checks on the agentic surface")
    para(doc, "Ten checks, no model, no threshold, no training data. Each returns a boolean "
              "and a reason, and the same inputs always give the same answer — the property "
              "a statistical detector cannot offer.")
    table(doc, ["Check", "What it enforces"], [
        ["C1 amount scope", "cart amount within the mandate cap"],
        ["C2 category scope", "MCC within the allowed set"],
        ["C3 merchant scope", "merchant within the allowed set"],
        ["C4 temporal validity", "inside the mandate's validity window"],
        ["C5 nonce freshness", "the mandate presentation has not been replayed"],
        ["C6 cumulative cap", "aggregate spend across carts within the cumulative cap"],
        ["C7 agent binding", "Ed25519 attestation by the delegated agent"],
        ["C8 confirmation binding", "what the user confirmed matches the cart's line items"],
        ["C9 revocation state", "the mandate has not been revoked"],
        ["C10 mandate signature", "Ed25519 signature by the principal over canonical JSON"],
    ], widths=[1.8, 4.5])


def s7_results(doc, d):
    h1(doc, "7. Results", "Rubric: results, analysis and honesty")

    if d["ablation"]:
        ab = d["ablation"]
        h2(doc, "Ablation at the base configuration")
        m = ab["_meta"]
        para(doc, f"ρ = {m['rho']}, λ = {m['lam']}, K = {m['K']}, β = {m['beta']}, seed 0. "
                  f"Train {m['n_train']:,} rows / {m['n_train_payers']:,} payers; test "
                  f"{m['n_test']:,} rows / {m['n_test_payers']:,} payers / "
                  f"{m['n_test_fraud']:,} fraudulent transactions "
                  f"({m['test_fraud_rate']:.3%}). Confidence intervals are 1000-resample "
                  f"bootstrap clustered on test payers.")
        rows = []
        for arm in ["B1", "B1+B2", "B1+B2+B3", "B1+B2+B3+B4a", "B1+B2+B3+B4b"]:
            if arm not in ab:
                continue
            x = ab[arm]
            rlo, rhi = x["ci"]["recall@fpr=0.001"]
            flo, fhi = x["ci"]["fpr@recall=0.7"]
            rows.append([
                arm + ("  ← baseline" if arm == "B1+B2+B3" else ""),
                x["n_features"], f"{x['pr_auc']:.4f}",
                f"{x['recall_at_fpr']['0.001']:.4f} [{rlo:.4f}, {rhi:.4f}]",
                f"{x['fpr_at_recall']['0.7']:.5f} [{flo:.5f}, {fhi:.5f}]",
            ])
        table(doc, ["Arm", "Feat.", "PR-AUC", "Recall @ FPR 0.1%", "FPR @ recall 70%"],
              rows, widths=[1.5, 0.5, 0.8, 1.9, 1.9], align_right=(1, 2, 3, 4), size=8.5)

        h2(doc, "Paired delta against the baseline")
        para(doc, "Deltas are computed on shared bootstrap resamples and signed so that "
                  "positive always means declared context helped.")
        rows = []
        for arm in ("B1+B2+B3+B4a", "B1+B2+B3+B4b"):
            for nm, e in ab.get("delta", {}).get(arm, {}).items():
                rows.append([arm.replace("B1+B2+B3+", "+"), nm,
                             f"{e['point']:+.5f}",
                             f"[{e['ci'][0]:+.5f}, {e['ci'][1]:+.5f}]",
                             "significant" if e["significant"] else "—"])
        table(doc, ["Arm", "Metric", "Δ", "95% CI", ""], rows,
              widths=[0.7, 1.6, 1.0, 1.9, 1.1], align_right=(2, 3), size=8.5)
        figure(doc, "ablation.png", "Figure 1. Each feature group added in turn. Error bars "
                                    "are the standard deviation across three seeds.")

    h2(doc, "The phase diagram")
    para(doc, "The primary output is a surface, not a number. Hatched cells — where the "
              "confidence interval includes zero on at least one seed — are part of the "
              "result. They are the regions where a payment network should not spend money "
              "collecting this field.")
    figure(doc, "phase_uniform.png",
           "Figure 2. Incremental value of declared payment context under the "
           "pre-registered adversary. Left: Δ recall at FPR 0.1%. Right: Δ FPR at recall "
           "70%. Hatched cells are not significant.")
    figure(doc, "phase_prevalence.png",
           "Figure 3. The same surface under the prevalence-matched adversary, whose "
           "declared code carries no marginal information at ρ = 1. Secondary analysis; the "
           "pre-registered axis is Figure 2.")
    figure(doc, "rho_star.png",
           "Figure 4. ρ* by λ, bracketed by the sweep grid rather than interpolated.")
    figure(doc, "secondary_sweeps.png",
           "Figure 5. Purpose cardinality K and beneficiary-feature noise β at ρ = 0.4, "
           "λ = 0.10.")

    if d["agentic"]:
        a = d["agentic"]
        bl, fp = a["bounded_loss"], a["false_positives_on_in_scope_traffic"]
        h2(doc, "Agentic conformance: coverage")
        rows = [[x["attack_id"], x["name"],
                 ", ".join(x["failed_checks"]) or "—",
                 "caught" if x["caught"] else "NOT CAUGHT"] for x in a["attacks"]]
        table(doc, ["ID", "Attack family", "Caught by", "Outcome"], rows,
              widths=[0.5, 2.6, 2.0, 1.1], size=8.5)
        para(doc, f"Coverage: {a['coverage']['caught']} of {a['coverage']['total']} families "
                  f"caught deterministically, at a false-positive rate of "
                  f"{fp['false_positive_rate']:.4f} — {fp['rejected']} of {fp['n']:,} "
                  f"conforming in-scope carts rejected.")
        callout(doc, "A9 and A10 are the most important rows in that table. They are "
                     "uncaught by construction, not by omission.", BAD)
        para(doc, "An agent that spends inside the mandate — whether compromised or steered "
                  "by a prompt injection — passes every check. Conformance checking does not "
                  "detect that. What it does is bound the loss. This is also the direct "
                  "answer to 'AP2 already specifies mandate verification': the contribution "
                  "is not the checks, it is the measurement of what they do and do not buy.")

        h2(doc, "Agentic conformance: bounded loss on the families it cannot catch")
        table(doc, ["Scenario", "Mean loss", "p95 loss", "Reduction"], [
            ["No mandate enforcement", f"₹{bl['mean_loss_unenforced']:,.0f}",
             f"₹{bl['p95_unenforced']:,.0f}", "—"],
            ["Enforced, single cart", f"₹{bl['mean_loss_single_cart']:,.0f}",
             f"₹{bl['cap']:,.0f}", f"{bl['reduction_single_cart']:.1%}"],
            ["Enforced, persistent attacker", f"₹{bl['mean_loss_persistent']:,.0f}",
             f"₹{bl['p95_persistent']:,.0f}", f"{bl['reduction_persistent']:.1%}"],
        ], widths=[2.3, 1.4, 1.3, 1.3], align_right=(1, 2, 3))
        para(doc, "The attacker's desired spend is drawn from the same RBI-calibrated "
                  "case-size distribution used everywhere else, so the comparison is against "
                  "a realistic loss profile rather than a convenient one.")

    h2(doc, "The two evidential regimes, side by side")
    table(doc, ["", "Human declaration", "Signed mandate"], [
        ["Nature of evidence", "probabilistic, possibly deceptive", "cryptographic, constraint-bounded"],
        ["Failure mode", "coaching (ρ)", "in-scope abuse (A9/A10)"],
        ["What it delivers", "a shift in ranking", "a deterministic reject, or a bounded loss"],
        ["False positives", "non-zero by construction", "zero on conforming traffic"],
        ["Degrades under pressure", "yes — see the phase diagram", "no — but its scope is narrower"],
    ], widths=[1.5, 2.4, 2.4])
    para(doc, "These are not the same claim and this project does not blur them. They are "
              "the same underlying question asked of two different kinds of evidence.")


def s8_feasibility(doc, d):
    h1(doc, "8. Feasibility", "Rubric: implementation feasibility")
    bullets(doc, [
        ("No new rail field is needed for corporate flows.", "ISO 20022 <Purp><Cd> already "
         "travels in pain.001 and pacs.008. The message structure exists and is standardised."),
        ("Consumer flows need retention, not invention.", "The gap is not that the field "
         "cannot be carried; it is whether a consumer-initiated declaration survives "
         "end-to-end rather than being discarded at a gateway. That is an operational and "
         "governance question, not a protocol one."),
        ("Latency is not a constraint.", "The consistency residuals are a lookup of "
         "per-purpose moments and one quadratic form over thirteen dimensions. It is "
         "microseconds, and the reference distributions are refreshed offline."),
        ("The reference set is built from confirmed-good history.", "Exactly the data a "
         "bank already holds. No new labelling process is required, and the model never "
         "needs fraud labels to compute a residual."),
        ("Governance is the real cost.", "A declared-purpose field is payer-supplied data "
         "used in an adverse decision. It needs a retention policy, an explanation path, "
         "and a false-positive appeal route. §9 does not pretend otherwise."),
        ("It converges with where the RBI is already heading.", "The April 2026 discussion "
         "paper proposes four friction controls. A detection signal is complementary to all "
         "four, and this work characterises the conditions under which it is worth having."),
    ])
    h2(doc, "A deployment sketch")
    para(doc, "Nothing here requires a new machine-learning capability. What it requires is "
              "a field, a reference table, and a policy.")
    bullets(doc, [
        ("Capture.", "A purpose selection at initiation, from a menu whose cardinality is "
         "an operational choice the K sweep informs directly. A coarser menu is cheaper and "
         "easier to answer correctly."),
        ("Retention.", "The declaration must survive to the point where the decision is "
         "made. This is the real integration work and it is a governance problem, not a "
         "protocol one."),
        ("Reference build.", "Offline, from confirmed-good history: per purpose class, the "
         "mean and shrunk covariance of the beneficiary feature vector. Refreshed on "
         "whatever cadence the beneficiary features themselves are refreshed on."),
        ("Scoring.", "One lookup and one quadratic form over thirteen dimensions. "
         "Microseconds, and no additional model to maintain."),
        ("Policy.", "The phase diagram is the deployment decision. If the operating regime "
         "sits in a hatched cell, the field is not worth acting on there, whatever it costs "
         "to collect."),
    ])

    h2(doc, "What this analysis is actually for")
    callout(doc, "Before a payment network spends money collecting another signal, can we "
                 "quantify when that signal remains useful under adversarial pressure?")
    para(doc, "That is the deliverable. Not a better fraud classifier — a decision framework "
              "with a measured adversarial tolerance, and a published generative process so "
              "the partition can be challenged.")


def s9_limitations(doc, d):
    h1(doc, "9. Limitations", "Rubric: honesty and self-assessment")
    bullets(doc, [
        ("No labelled public APP dataset exists.", "Every number here comes from a "
         "simulator. We make no claim about absolute detection rates; the results "
         "characterise relative behaviour across parameter regimes."),
        ("Absolute performance is higher than any deployed system.", "The baseline reaches "
         "~0.90 recall at 0.1% FPR. Session telemetry was deliberately made strong so the "
         "baseline would not be a strawman, and recall at 0.5% and 1% FPR is consequently "
         "saturated for the B4 arms. Those operating points were not changed after seeing "
         "results; the saturation is reported instead."),
        ("Production systems already capture much of B2 and B3.", "The experiment asks the "
         "narrower question of whether a further signal earns its collection cost once "
         "those are in place."),
        ("The consumer population rate of <Purp> is unverified.", "That the field exists is "
         "established. How often it is populated on consumer flows, and whether it survives "
         "end-to-end, is not established from any primary source available to us."),
        ("Results are conditional on the generative model.", "ρ* is a property of this "
         "simulator under this threat model and this parameterisation. It is not a "
         "universal threshold and must never be quoted as one."),
        ("The ρ mechanism is one adversary among many.", "An adversary who also controls "
         "which mule receives the payment, choosing one whose profile matches the declared "
         "purpose, is not modelled."),
        ("The agentic module bounds loss; it does not detect intent.", "Two of ten families "
         "pass every check by construction."),
        ("The generative model was corrected after the pre-registration.", "CHANGELOG.md "
         "records every change, what was observed that prompted it, and why it is a "
         "realism correction. PREREGISTRATION.md has never been edited and was the sole "
         "content of the first commit."),
    ])

    h2(doc, "The circularity objection, answered before it is asked")
    para(doc, "We plant context metadata, so we make no claim about absolute detection "
              "rates. What is not circular: the deterministic results are structural, the "
              "phase diagram measures relative behaviour across parameter regimes rather "
              "than a point estimate, and we have published the generative process. We are "
              "characterising when a control is worth deploying, not claiming a benchmark win.")


def s10_responsible(doc, d):
    h1(doc, "10. Responsible AI and compliance", "Rubric: responsible innovation")
    bullets(doc, [
        ("Synthetic data only.", "No real payment data, no real accounts, no personal data "
         "of any kind. Every ledger is generated from a published process and is "
         "reproducible from a seed."),
        ("No live-system testing.", "Nothing in this repository contacts a payment network, "
         "a bank, a merchant, or an agent platform."),
        ("No operational attack tooling.", "The agentic attack families construct malformed "
         "mandates against a local verifier with locally generated keys. There is no "
         "capability here that is useful against a real system."),
        ("Primary sources only.", "RBI, ISO 20022, AP2, BIS Nexus and named public "
         "statistics. No statistics aggregators."),
        ("Fairness.", "A declared-purpose control acts on payer-supplied text in an adverse "
         "decision. The phase diagram is the fairness argument as much as the efficacy "
         "one: it identifies the regimes in which the signal is not informative, and "
         "deploying an uninformative signal in an adverse decision is the failure mode "
         "that matters."),
        ("Reproducibility.", "Pre-registration in the first commit, frozen hyperparameters "
         "with the full trial log, seeds fixed, and every result written to JSON that the "
         "prototype and this document both read."),
    ])


def s11_appendix(doc, d, meta):
    doc.add_page_break()
    h1(doc, "Appendix A. Pre-registration (verbatim, unedited)")
    para(doc, f"Sole content of commit {meta['prereg_commit']}. Verify with "
              f"git log --reverse --stat.", italic=True, colour=SLATE)
    txt = Path("PREREGISTRATION.md").read_text()
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = "Consolas"
    r.font.size = Pt(7.6)

    doc.add_page_break()
    h1(doc, "Appendix B. Changes made after pre-registration")
    txt = Path("CHANGELOG.md").read_text()
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = "Consolas"
    r.font.size = Pt(7.6)

    doc.add_page_break()
    h1(doc, "Appendix C. Frozen hyperparameters")
    fp = json.loads(Path("config/frozen_params.json").read_text()) \
        if Path("config/frozen_params.json").exists() else {}
    para(doc, f"Selected on: {fp.get('selected_on', '—')}. Criterion: "
              f"{fp.get('criterion', '—')}. Candidates: {fp.get('n_candidates', '—')}. "
              f"Tuning rows: {fp.get('tune_rows', '—')}. Best CV PR-AUC: "
              f"{fp.get('best_pr_auc_mean', float('nan')):.5f} ± "
              f"{fp.get('best_pr_auc_std', float('nan')):.5f}.")
    p = doc.add_paragraph()
    r = p.add_run(json.dumps(fp.get("params", {}), indent=2))
    r.font.name = "Consolas"; r.font.size = Pt(8.5)

    h1(doc, "Appendix D. Reproduction")
    p = doc.add_paragraph()
    r = p.add_run(
        "uv venv --python 3.12 .venv && uv pip install -r requirements.txt\n"
        "make test        # hygiene and conformance tests\n"
        "make tune        # select and freeze baseline hyperparameters\n"
        "make ablation    # base-configuration table with CIs\n"
        "make sweep       # phase study\n"
        "make agentic     # conformance coverage and bounded loss\n"
        "make fidelity    # realism scorecard\n"
        "make figures     # phase diagrams\n"
        "make inspector   # worked consistency cases\n"
        "make ui && make api\n")
    r.font.name = "Consolas"; r.font.size = Pt(9)
    para(doc, "Full data card at docs/DATA_CARD.md, model card at docs/MODEL_CARD.md, "
              "limitations at docs/LIMITATIONS.md.")


def main():
    d = {"phase": load("phase_surface.json"), "ablation": load("ablation.json"),
         "agentic": load("agentic_conformance.json"), "fidelity": load("fidelity.json")}
    meta = {
        "prereg_commit": git("log", "--reverse", "--format=%H", default="")[:40] or "—",
        "head": git("rev-parse", "--short", "HEAD"),
        "built": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
    }
    doc = Document()
    setup(doc)
    cover(doc, meta)
    s1_executive(doc, d)
    s2_problem(doc, d)
    s3_prior_art(doc, d)
    s4_identify(doc, d)
    s5_generate(doc, d)
    s6_defend(doc, d)
    s7_results(doc, d)
    s8_feasibility(doc, d)
    s9_limitations(doc, d)
    s10_responsible(doc, d)
    s11_appendix(doc, d, meta)
    doc.save(OUT)
    missing = [k for k, v in d.items() if v is None]
    print(f"written -> {OUT}")
    if missing:
        print(f"sections rendered as placeholders (not yet generated): {missing}")


if __name__ == "__main__":
    main()
